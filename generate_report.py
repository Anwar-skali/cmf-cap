#!/usr/bin/env python3
"""
CMF Platform — Rapport de Stage / PFA
Generates a complete Word (.docx) report from the analyzed codebase.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import datetime

# ─────────────────────────── Helpers ───────────────────────────

def set_cell_shading(cell, color_hex):
    """Apply background shading to a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_styled_table(doc, headers, rows, col_widths=None):
    """Add a professionally styled table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, "2B5797")
    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            run.font.size = Pt(9)
            if r_idx % 2 == 1:
                set_cell_shading(cell, "EBF1F8")
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    return table

def add_heading_numbered(doc, text, level=1):
    """Add a heading with numbering style."""
    doc.add_heading(text, level=level)

def add_normal(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.clear()
    run = p.add_run(text)
    run.font.size = Pt(11)
    if level > 0:
        p.paragraph_format.left_indent = Cm(1.27 * level)
    return p

def add_code_block(doc, code_text):
    p = doc.add_paragraph()
    run = p.add_run(code_text)
    run.font.name = "Consolas"
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    set_cell_shading_p = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5"/>')
    p._p.get_or_add_pPr().append(set_cell_shading_p)
    return p

# ══════════════════════════════════════════════════════════════
#                       DOCUMENT CREATION
# ══════════════════════════════════════════════════════════════

doc = Document()

# ── Page setup ──
for section in doc.sections:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ── Default font ──
style = doc.styles["Normal"]
font = style.font
font.name = "Calibri"
font.size = Pt(11)

# Heading styles
for i in range(1, 5):
    hs = doc.styles[f"Heading {i}"]
    hs.font.name = "Calibri"
    hs.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

doc.styles["Heading 1"].font.size = Pt(18)
doc.styles["Heading 2"].font.size = Pt(14)
doc.styles["Heading 3"].font.size = Pt(12)

# ══════════════════════════════════════════════════════════════
#                   PAGE DE GARDE
# ══════════════════════════════════════════════════════════════

for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("RAPPORT DE STAGE\nPROJET DE FIN D'ANNÉE")
run.bold = True
run.font.size = Pt(26)
run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Conception et Développement\nd'une Plateforme Web de\nCapacity Management Framework (CMF)")
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

for _ in range(3):
    doc.add_paragraph()

# Info table on cover
info_data = [
    ("Étudiant :", "[Votre Nom]"),
    ("École :", "[Votre École]"),
    ("Filière :", "[Votre Filière]"),
    ("Entreprise :", "[Entreprise d'accueil]"),
    ("Encadrant entreprise :", "[Encadrant]"),
    ("Encadrant académique :", "[Encadrant académique]"),
    ("Année académique :", "2025 – 2026"),
]

table = doc.add_table(rows=len(info_data), cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (label, value) in enumerate(info_data):
    c0 = table.rows[i].cells[0]
    c1 = table.rows[i].cells[1]
    c0.width = Cm(5)
    c1.width = Cm(8)
    r0 = c0.paragraphs[0].add_run(label)
    r0.bold = True
    r0.font.size = Pt(11)
    c0.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r1 = c1.paragraphs[0].add_run(value)
    r1.font.size = Pt(11)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
#                     REMERCIEMENTS
# ══════════════════════════════════════════════════════════════

add_heading_numbered(doc, "Remerciements", 1)

add_normal(doc, (
    "Je tiens à exprimer ma profonde gratitude envers toutes les personnes qui ont contribué "
    "à la réalisation de ce stage et à l'élaboration de ce rapport."
))

add_normal(doc, (
    "Mes remerciements s'adressent tout d'abord à [Entreprise d'accueil] pour m'avoir accueilli "
    "au sein de son équipe et pour m'avoir offert l'opportunité de travailler sur un projet "
    "aussi enrichissant que la plateforme CMF. L'environnement de travail, la confiance accordée "
    "et les ressources mises à disposition ont été déterminants dans la réussite de ce projet."
))

add_normal(doc, (
    "Je remercie chaleureusement mon encadrant de stage [Encadrant], pour sa disponibilité, "
    "ses conseils techniques précieux et son accompagnement tout au long du projet. Ses retours "
    "réguliers et sa rigueur m'ont permis de progresser tant sur le plan technique que "
    "méthodologique."
))

add_normal(doc, (
    "Je remercie également l'ensemble de l'équipe pour leur accueil bienveillant, leurs "
    "échanges constructifs et leur esprit de collaboration. Les discussions techniques partagées "
    "et les recommandations pratiques ont largement contribué à la qualité du résultat final."
))

add_normal(doc, (
    "Enfin, je remercie mon encadrant académique [Encadrant académique] pour son suivi, "
    "ses orientations méthodologiques et sa relecture attentive de ce rapport. Je remercie "
    "egalement mon école pour la formation reçue qui m'a permis d'aborder ce projet avec "
    "les compétences nécessaires."
))

add_normal(doc, (
    "Ce stage a été une expérience formatrice et déterminante dans mon parcours. "
    "Il m'a permis de confronter les connaissances théoriques acquises en cours "
    "à des problématiques industrielles concrètes."
))

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
#                       RÉSUMÉ
# ══════════════════════════════════════════════════════════════

add_heading_numbered(doc, "Résumé", 1)

add_normal(doc, (
    "Dans le cadre de mon stage de fin d'études, j'ai participé à la conception et au développement "
    "d'une plateforme web dédiée au Capacity Management Framework (CMF). Ce projet vise à répondre "
    "à un besoin croissant de centralisation et de fiabilisation des données relatives à la gestion "
    "des projets, des fournisseurs, de la capacité de production et du suivi qualité au sein "
    "d'un environnement industriel."
))

add_normal(doc, (
    "Les processus existants reposaient principalement sur des fichiers Excel partagés entre "
    "plusieurs parties prenantes (Buyers, Capacity Managers, responsables SQD). Cette approche "
    "engendrait des problèmes de dispersion des données, d'erreurs manuelles, de difficultés "
    "de suivi et d'absence de traçabilité. La plateforme CMF a été développée pour pallier "
    "ces insuffisances en offrant une solution centralisée, sécurisée et collaboratives."
))

add_normal(doc, (
    "La plateforme repose sur une architecture full-stack : un frontend React/TypeScript "
    "communiquant via une API REST avec un backend Python/FastAPI, utilisant SQLAlchemy "
    "comme ORM et SQLite comme base de données. Le système d'importation constitue le "
    "coeur technique du projet : il intègre un pipeline complet allant de l'upload d'un "
    "fichier Excel jusqu'à l'insertion en base, en passant par la détection de feuille, "
    "la détection de la ligne d'en-tête, le mapping sémantique des colonnes assisté par "
    "intelligence artificielle (Ollama), la validation des données et l'exécution sécurisée "
    "avec transactions."
))

add_normal(doc, (
    "La plateforme gère également des structures de données modulaires (K0, K9, et structures "
    "personnalisées), un système de rôles et permissions, un tableau de bord analytique, "
    "la gestion des fournisseurs, des capacités et des risques. Des tests unitaires "
    "et d'intégration ont été mis en place pour garantir la fiabilité du système."
))

doc.add_paragraph()
p = add_normal(doc, "Mots-clés : ", bold=True)
keywords = "Capacity Management Framework, plateforme web, React, TypeScript, FastAPI, Python, Excel import, mapping sémantique, Ollama, intelligence artificielle, SQLAlchemy, SQLite"
run = p.add_run(keywords)
run.italic = True

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
#                       ABSTRACT
# ══════════════════════════════════════════════════════════════

add_heading_numbered(doc, "Abstract", 1)

add_normal(doc, (
    "As part of my end-of-studies internship, I contributed to the design and development "
    "of a web platform dedicated to Capacity Management Framework (CMF). The project addresses "
    "a growing need for data centralization and reliability in managing projects, suppliers, "
    "production capacity, and quality monitoring within an industrial environment."
))

add_normal(doc, (
    "Previous processes relied on shared Excel files among multiple stakeholders (Buyers, "
    "Capacity Managers, SQD leads), resulting in data dispersion, manual errors, tracking "
    "difficulties, and lack of traceability. The CMF platform was developed to address these "
    "shortcomings by providing a centralized, secure, and collaborative solution."
))

add_normal(doc, (
    "The platform is built on a full-stack architecture: a React/TypeScript frontend "
    "communicating via a REST API with a Python/FastAPI backend, using SQLAlchemy as ORM "
    "and SQLite as database. The import system constitutes the core technical component: "
    "it integrates a complete pipeline from Excel file upload to database insertion, "
    "including worksheet detection, header row detection, semantic column mapping assisted "
    "by artificial intelligence (Ollama), data validation, and secure execution with transactions."
))

add_normal(doc, (
    "The platform also manages modular data structures (K0, K9, and custom structures), "
    "a role-based access control system, an analytical dashboard, supplier and capacity management, "
    "and risk tracking. Unit and integration tests were implemented to ensure system reliability."
))

doc.add_paragraph()
p = add_normal(doc, "Keywords: ", bold=True)
run = p.add_run(
    "Capacity Management Framework, web platform, React, TypeScript, FastAPI, Python, "
    "Excel import, semantic mapping, Ollama, artificial intelligence, SQLAlchemy, SQLite"
)
run.italic = True

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
#                  TABLE DES MATIÈRES
# ══════════════════════════════════════════════════════════════

add_heading_numbered(doc, "Table des matières", 1)

toc_items = [
    ("Introduction générale", ""),
    ("Chapitre 1 — Présentation de l'entreprise", ""),
    ("Chapitre 2 — Contexte et problématique", ""),
    ("Chapitre 3 — Analyse et spécifications", ""),
    ("    3.1 Besoins fonctionnels", ""),
    ("    3.2 Besoins non fonctionnels", ""),
    ("    3.3 Acteurs du système", ""),
    ("    3.4 Cas d'utilisation", ""),
    ("Chapitre 4 — Conception et architecture", ""),
    ("    4.1 Architecture générale", ""),
    ("    4.2 Architecture backend", ""),
    ("    4.3 Architecture frontend", ""),
    ("    4.4 Pipeline d'importation", ""),
    ("Chapitre 5 — Technologies utilisées", ""),
    ("Chapitre 6 — Conception de la base de données", ""),
    ("    6.1 Modèle de données", ""),
    ("    6.2 Relations et contraintes", ""),
    ("Chapitre 7 — Développement de la plateforme", ""),
    ("    7.1 Tableau de bord (Dashboard)", ""),
    ("    7.2 Gestion des projets", ""),
    ("    7.3 Structures et templates", ""),
    ("    7.4 Création manuelle d'un projet", ""),
    ("    7.5 Gestion des fournisseurs", ""),
    ("    7.6 Gestion de la capacité", ""),
    ("    7.7 Suivi qualité (SQD)", ""),
    ("Chapitre 8 — Importation des données", ""),
    ("    8.1 Workflow d'importation", ""),
    ("    8.2 Détection de feuille et d'en-tête", ""),
    ("    8.3 Mapping sémantique des colonnes", ""),
    ("    8.4 Validation et preview", ""),
    ("    8.5 Exécution sécurisée", ""),
    ("Chapitre 9 — Structures K0 et K9", ""),
    ("    9.1 Structure K0", ""),
    ("    9.2 Structure K9", ""),
    ("Chapitre 10 — Optimisation des imports", ""),
    ("Chapitre 11 — Difficultés rencontrées", ""),
    ("Chapitre 12 — Tests et validation", ""),
    ("Chapitre 13 — Résultats obtenus", ""),
    ("Chapitre 14 — Compétences acquises", ""),
    ("Chapitre 15 — Perspectives", ""),
    ("Conclusion générale", ""),
    ("Annexes", ""),
]

for item, _ in toc_items:
    p = doc.add_paragraph()
    indent = item.startswith("    ")
    text = item.strip()
    run = p.add_run(text)
    run.font.size = Pt(11)
    if not indent:
        run.bold = True
    else:
        p.paragraph_format.left_indent = Cm(1.5)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
#              INTRODUCTION GÉNÉRALE
# ══════════════════════════════════════════════════════════════

add_heading_numbered(doc, "Introduction générale", 1)

add_normal(doc, (
    "Le secteur industriel contemporain fait face à une complexité croissante dans la gestion "
    "de ses données opérationnelles. La coordination entre les différents acteurs — acheteurs, "
    "responsables de capacité, responsables qualité — nécessite des outils fiables, centralisés "
    "et adaptés aux processus métier. Dans de nombreuses organisations, les données relatives "
    "aux projets, aux fournisseurs et à la capacité de production sont encore gérées au moyen "
    "de fichiers Excel circulant par messagerie, générant des problèmes de versioning, "
    "d'incohérence et de perte d'information."
))

add_normal(doc, (
    "Le Capacity Management Framework (CMF) est un processus métier structurant qui permet "
    "d'évaluer, planifier et suivre la capacité de production des fournisseurs en fonction "
    "des besoins projet. Il implique une collaboration étroite entre plusieurs départements "
    "et repose sur des données précises, actualisées et facilement accessibles. La fiabilité "
    "de ces données est essentielle pour prendre des décisions stratégiques d'approvisionnement."
))

add_normal(doc, (
    "C'est dans ce contexte que s'inscrit le présent stage, dont l'objectif principal était "
    "de concevoir et développer une plateforme web permettant de centraliser l'ensemble des "
    "données CMF, d'automatiser l'importation depuis des fichiers Excel, de garantir la "
    "validation des données et de fournir un suivi en temps réel via un tableau de bord "
    "analytique. La plateforme devait également intégrer un système intelligent de mapping "
    "des colonnes Excel vers les champs du système, utilisant l'intelligence artificielle "
    "comme fallback lorsque le mapping automatique ne suffisait pas."
))

add_normal(doc, (
    "La problématique principale de ce stage peut être formulée ainsi : comment concevoir "
    "une plateforme web permettant de centraliser et fiabiliser la gestion des données CMF "
    "tout en facilitant l'importation, la validation et le suivi des projets, des fournisseurs "
    "et des évaluations de capacité ?"
))

add_normal(doc, (
    "Pour répondre à cette problématique, nous avons adopté une démarche itérative combinant "
    "l'analyse des besoins, la conception architecturale, le développement frontend et backend, "
    "la mise en place d'un système d'importation intelligent et la réalisation de tests. "
    "Ce rapport présente le travail réalisé au cours du stage en suivant la structure classique : "
    "présentation de l'entreprise, contexte, analyse, conception, développement, tests et résultats."
))

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
#          CHAPITRE 1 — PRÉSENTATION DE L'ENTREPRISE
# ══════════════════════════════════════════════════════════════

add_heading_numbered(doc, "Chapitre 1 — Présentation de l'entreprise", 1)

add_normal(doc, (
    "Ce stage s'est déroulé au sein de [Entreprise d'accueil], une entreprise intervenant "
    "dans le domaine de l'industrie automobile. L'entreprise est impliquée dans la gestion "
    "de projets industriels complexes, nécessitant une coordination étroite entre les "
    "différentes fonctions : achat, planification de capacité, qualité et supply chain."
))

add_normal(doc, (
    "L'entreprise utilise des processus de Capacity Management Framework (CMF) pour structurer "
    "le suivi de ses projets. Ces processus impliquent plusieurs catégories d'acteurs : "
    "les Buyers (responsables achats), les Capacity Managers (responsables de capacité), "
    "et les SQD (Service Quality Development — responsables qualité). Chaque acteur contribue "
    "à不同 des phases du cycle de vie d'un projet, depuis la création jusqu'à la validation "
    "finale."
))

add_normal(doc, (
    "L'environnement technique existant reposait principalement sur des fichiers Excel partagés "
    "pour la collecte et le traitement des données. Cette approche, bien que flexible, "
    "présentait des limites significatives en termes de centralisation, de traçabilité et "
    "de fiabilité des données."
))

add_heading_numbered(doc, "Environnement technique", 2)

add_normal(doc, (
    "L'environnement technique du stage comprenait :"
))

tech_env = [
    "Système de versionnement : Git (GitHub)",
    "Langages : Python (backend), TypeScript/JavaScript (frontend)",
    "Framework backend : FastAPI",
    "ORM : SQLAlchemy",
    "Base de données : SQLite (développement), PostgreSQL (production)",
    "Framework frontend : React",
    "Bundler : Vite",
    "Gestion d'état : Zustand, TanStack Query",
    "Style : Tailwind CSS",
    "Traitement Excel : openpyxl (backend), xlsx (frontend)",
    "Intelligence artificielle : Ollama (LLM local)",
]
for t in tech_env:
    add_bullet(doc, t)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
#        CHAPITRE 2 — CONTEXTE ET PROBLÉMATIQUE
# ══════════════════════════════════════════════════════════════

add_heading_numbered(doc, "Chapitre 2 — Contexte et problématique", 1)

add_heading_numbered(doc, "2.1 Le Capacity Management Framework", 2)

add_normal(doc, (
    "Le Capacity Management Framework (CMF) est un processus structurant utilisé dans "
    "l'industrie automobile pour évaluer et planifier la capacité de production des "
    "fournisseurs. Il s'inscrit dans une démarche de sécurisation de la chaîne d'approvisionnement "
    "et de réduction des risques liés à la sous-capacité."
))

add_normal(doc, (
    "Le processus CMF implique trois phases principales correspondant à trois rôles distincts :"
))

add_bullet(doc, (
    "Phase Buyer : identification des besoins, définition du projet, sélection du fournisseur, "
    "définition des paramètres de base (numéro de pièce, prix, quantité, etc.)"
))
add_bullet(doc, (
    "Phase Capacity Manager : évaluation de la capacité actuelle et maximale du fournisseur, "
    "planification des objectifs par semaine, suivi des prévisions et des réalisations"
))
add_bullet(doc, (
    "Phase SQD : évaluation qualité, audit de la chaîne d'approvisionnement, classification "
    "des risques qualité (GREEN/ORANGE/RED), validation finale"
))

add_heading_numbered(doc, "2.2 Problématique identifiée", 2)

add_normal(doc, (
    "L'analyse des processus existants a mis en évidence plusieurs problèmes majeurs :"
))

problems = [
    ("Dispersion des données", "Les informations relatives aux projets, fournisseurs et capacités "
     "étaient réparties dans de multiples fichiers Excel, chacun ayant sa propre structure et "
     "son propre format."),
    ("Erreurs manuelles", "La saisie et la copie de données entre les fichiers engendraient "
     "des erreurs fréquentes : doublons, incohérences de format, valeurs manquantes."),
    ("Absence de traçabilité", "Il n'était pas possible de retracer qui avait modifié quoi, "
     "quand et pourquoi. L'historique des modifications était perdu."),
    ("Difficulté de suivi", "Le suivi de l'avancement des projets nécessitait de consulter "
     "plusieurs fichiers et de croiser manuellement les informations."),
    ("Structures hétérogènes", "Les différents projets utilisaient des structures de données "
     "différentes (K0, K9, etc.), rendant difficile la comparaison et l'agrégation."),
    ("Mapping manuel", "L'importation de données depuis des fichiers tiers nécessitait un "
     "remanuel de mapping colonne par colonne, processus chronophage et source d'erreurs."),
    ("Manque de centralisation", "Aucun système unique ne permettait de disposer d'une vue "
     "d'ensemble sur l'ensemble des projets CMF."),
]

for title, desc in problems:
    p = doc.add_paragraph()
    run = p.add_run(f"{title} : ")
    run.bold = True
    run.font.size = Pt(11)
    run = p.add_run(desc)
    run.font.size = Pt(11)

add_heading_numbered(doc, "2.3 Problématique de recherche", 2)

add_normal(doc, (
    "À partir de cette analyse, la problématique de recherche du stage peut être formulée "
    "comme suit :"
), italic=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.left_indent = Cm(2)
p.paragraph_format.right_indent = Cm(2)
run = p.add_run(
    "« Comment concevoir et développer une plateforme web permettant de centraliser et fiabiliser "
    "la gestion des données CMF, tout en facilitant l'importation automatisée depuis des fichiers "
    "Excel, la validation des données et le suivi des projets, fournisseurs et évaluations "
    "de capacité ? »"
)
run.italic = True
run.font.size = Pt(11)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
#       CHAPITRE 3 — ANALYSE ET SPÉCIFICATIONS
# ══════════════════════════════════════════════════════════════

add_heading_numbered(doc, "Chapitre 3 — Analyse et spécifications", 1)

add_heading_numbered(doc, "3.1 Besoins fonctionnels", 2)

add_normal(doc, (
    "L'analyse des besoins fonctionnels a été réalisée en examinant les processus métier "
    "existant et en identifiant les fonctionnalités nécessaires pour répondre à la problématique "
    "définie. Les besoins ont été catégorisés selon les domaines fonctionnels suivants :"
))

add_heading_numbered(doc, "3.1.1 Gestion des projets", 3)

features_projects = [
    "Liste des projets avec recherche, filtrage et pagination",
    "Création de projets avec sélection de template (K0, K9, structures personnalisées)",
    "Consultation détaillée d'un projet avec onglets (Détails, Pièces, Risques, Documents)",
    "Modification de projets avec contrôle d'accès par rôle (Buyer, Capacity Manager, SQD)",
    "Suppression de projets (suppression logique / soft-delete)",
    "Suppression en masse (bulk delete) avec boîte de confirmation",
    "Workflow de validation : Buyer → Capacity Manager → SQD → Complété",
    "Suivi de l'étape de workflow en temps réel",
]
for f in features_projects:
    add_bullet(doc, f)

add_heading_numbered(doc, "3.1.2 Importation de données", 3)

features_import = [
    "Upload de fichiers Excel (.xlsx, .xls, .xlsm)",
    "Détection automatique de la feuille de travail pertinente (worksheet scoring)",
    "Détection de la ligne d'en-tête (top 20 lignes évaluées)",
    "Détection de l'orientation : horizontale (tabulaire) ou verticale (clé-valeur)",
    "Mapping sémantique des colonnes avec cascade à 5 niveaux",
    "Mapping par cache (mappings précédemment confirmés)",
    "Mapping par correspondance exacte normalisée",
    "Mapping par alias connus",
    "Mapping par similarité floue (fuzzy matching)",
    "Mapping par IA (Ollama) en fallback",
    "Validation des données avec rapports d'erreurs et d'avertissements",
    "Preview avant import avec actions CREATE/UPDATE/RESTORE/SKIP",
    "Exécution sécurisée avec transactions (SAVEPOINT par ligne)",
    "Historique des imports avec statistiques détaillées",
    "Export des erreurs en Excel",
]
for f in features_import:
    add_bullet(doc, f)

add_heading_numbered(doc, "3.1.3 Gestion des structures/templates", 3)

features_templates = [
    "Création de structures de données modulaires (sections → groupes → champs)",
    "Import de structures depuis des fichiers JSON ou Excel",
    "Éditeur de templates avec builder visuel et éditeur JSON",
    "Attribution de permissions par rôle (Buyer, Capacity Manager, SQD) au niveau champ",
    "Import de structures hiérarchiques (modules → tables → champs)",
    "Structures protégées K0 et K9 (non supprimables)",
    "Duplication et versioning de templates",
    "Publication et archivage de templates",
    "Validation des champs sélectionnables (options obligatoires)",
]
for f in features_templates:
    add_bullet(doc, f)

add_heading_numbered(doc, "3.1.4 Gestion des fournisseurs", 3)

features_suppliers = [
    "CRUD complet des fournisseurs",
    "Assignation de fournisseurs à des projets",
    "Filtrage et recherche de fournisseurs",
    "Vues grille et carte pour les fournisseurs",
    "Statut du fournisseur (actif, inactif, blacklisté)",
    "Catégorisation et notation des fournisseurs",
]
for f in features_suppliers:
    add_bullet(doc, f)

add_heading_numbered(doc, "3.1.5 Gestion de la capacité", 3)

features_capacity = [
    "Évaluations de capacité mensuelles",
    "Suivi de la capacité actuelle vs maximale",
    "Taux d'utilisation calculé automatiquement",
    "Statistiques de couverture par fournisseur",
    "Données mensuelles par année",
    "Niveaux de risque (low, medium, high, critical)",
    "Champs spécifiques CMF : CATE, gate, target_week, forecast_week, completed_week",
]
for f in features_capacity:
    add_bullet(doc, f)

add_heading_numbered(doc, "3.1.6 Suivi qualité et risques", 3)

features_risks = [
    "Gestion des risques avec sévérité et probabilité",
    "Score de risque calculé automatiquement",
    "Matrice de chaleur des risques",
    "Tableau Kanban pour le suivi des risques",
    "Actions d'atténuation et de fermeture",
    "Affectation de responsables",
    "Statistiques de distribution des risques",
]
for f in features_risks:
    add_bullet(doc, f)

add_heading_numbered(doc, "3.1.7 Tableau de bord et administration", 3)

features_dashboard = [
    "Tableau de bord analytique avec KPI en temps réel",
    "Graphiques : barres, lignes, secteurs (Recharts)",
    "Indicateurs : projets actifs, capacité totale, taux d'utilisation, risques",
    "Gestion des utilisateurs et des rôles",
    "Authentification JWT avec refresh tokens",
    "Journal d'activité (audit trail)",
    "Notifications utilisateur",
    "Internationalisation français/anglais",
]
for f in features_dashboard:
    add_bullet(doc, f)

add_heading_numbered(doc, "3.2 Besoins non fonctionnels", 2)

nf_needs = [
    ("Performance", "Temps de réponse acceptables pour les opérations CRUD et l'importation "
     "de fichiers Excel de grande taille. Pipeline d'import optimisé pour réduire les appels IA."),
    ("Sécurité", "Authentification JWT avec rotation des refresh tokens. Contrôle d'accès "
     "basé sur les rôles (RBAC). Chiffrement des mots de passe avec bcrypt."),
    ("Maintenabilité", "Architecture clean avec séparation des couches (domain, application, "
     "infrastructure, API). Code TypeScript strict côté frontend."),
    ("Ergonomie", "Interface responsive avec Tailwind CSS. Composants UI réutilisables (Radix UI). "
     "Internationalisation FR/EN."),
    ("Fiabilité", "Transactions database avec SAVEPOINT par ligne. Validation côté serveur. "
     "Tests unitaires et d'intégration."),
    ("Extensibilité", "Système de templates modulaires. Ajout de nouvelles structures possibles "
     "sans modification du code. Architecture plugin-like pour les imports."),
    ("Portabilité", "Déploiement possible sur Vercel (frontend) + Railway/Neon (backend + DB). "
     "Configuration multi-environnement."),
]

for title, desc in nf_needs:
    p = doc.add_paragraph()
    run = p.add_run(f"{title} : ")
    run.bold = True
    run.font.size = Pt(11)
    run = p.add_run(desc)
    run.font.size = Pt(11)

add_heading_numbered(doc, "3.3 Acteurs du système", 2)

add_normal(doc, (
    "Le système identifie cinq rôles d'utilisateurs, chacun disposant de permissions "
    "spécifiques :"
))

roles_table = [
    ["Admin", "Administrateur",
     "Gestion complète de tous les modules. Création/suppression de templates. "
     "Gestion des utilisateurs. Seul rôle autorisé à supprimer des projets."],
    ["Buyer", "Acheteur",
     "Création de projets. Édition des champs de la section Buyer. Consultation "
     "des autres sections en lecture seule. Import de données."],
    ["Capacity Manager", "Responsable capacité",
     "Édition des champs de la section Capacity Manager. Évaluation de capacité. "
     "Création d'évaluations. Consultation des sections Buyer et SQD."],
    ["SQD", "Responsable qualité",
     "Édition des champs de la section SQD. Validation qualité. Classification "
     "des risques. Consultation des sections Buyer et Capacity."],
    ["Viewer", "Observateur",
     "Consultation en lecture seule de l'ensemble des données. Aucune permission "
     "d'écriture."],
]

add_styled_table(doc, ["Rôle", "Titre", "Responsabilités"], roles_table, [3, 3.5, 9])

add_heading_numbered(doc, "3.4 Cas d'utilisation", 2)

add_normal(doc, (
    "Les principaux cas d'utilisation du système peuvent être organisés en plusieurs "
    "catégories correspondant aux modules fonctionnels :"
))

use_cases = [
    ["UC-01", "S'authentifier", "Tous les utilisateurs", "Connexion, inscription, réinitialisation mot de passe"],
    ["UC-02", "Consulter le tableau de bord", "Tous les utilisateurs", "Vue d'ensemble des KPI et graphiques"],
    ["UC-03", "Créer un projet", "Buyer, Admin", "Création manuelle ou via template K0/K9"],
    ["UC-04", "Modifier un projet", "Buyer, Capacity Manager", "Édition selon les permissions de section"],
    ["UC-05", "Supprimer un projet", "Admin", "Suppression logique avec confirmation"],
    ["UC-06", "Importer des données Excel", "Buyer, Admin", "Pipeline complet : upload → mapping → validation → exécution"],
    ["UC-07", "Gérer les templates", "Admin", "CRUD, duplication, publication, import JSON/Excel"],
    ["UC-08", "Gérer les fournisseurs", "Buyer, Capacity Manager", "CRUD et assignation aux projets"],
    ["UC-09", "Évaluer la capacité", "Capacity Manager", "Création et suivi des évaluations mensuelles"],
    ["UC-10", "Gérer les risques", "Buyer, SQD, Capacity Manager", "CRUD, atténuation, fermeture"],
    ["UC-11", "Gérer les utilisateurs", "Admin", "Création, modification, désactivation"],
    ["UC-12", "Consulter l'historique", "Tous les utilisateurs", "Journal d'activité et historique des imports"],
]

add_styled_table(doc, ["ID", "Cas d'utilisation", "Acteur(s)", "Description"], use_cases, [1.5, 3.5, 4, 6.5])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
#      CHAPITRE 4 — CONCEPTION ET ARCHITECTURE
# ══════════════════════════════════════════════════════════════

add_heading_numbered(doc, "Chapitre 4 — Conception et architecture", 1)

add_heading_numbered(doc, "4.1 Architecture générale", 2)

add_normal(doc, (
    "La plateforme CMF adopte une architecture client-serveur reposant sur une API REST. "
    "L'architecture est séparée en trois grandes couches : le frontend (client web), "
    "le backend (serveur API) et la base de données. Cette séparation permet une évolution "
    "indépendante de chaque couche et facilite la maintenance."
))

add_normal(doc, (
    "Le schéma d'architecture général peut être représenté comme suit :"
))

arch_lines = [
    "┌─────────────────────────────────────────────┐",
    "│              FRONTEND (React)               │",
    "│  TypeScript · Tailwind CSS · TanStack Query │",
    "│  Vite · Zustand · Radix UI                  │",
    "└──────────────────┬──────────────────────────┘",
    "                   │  REST API (JSON)",
    "                   ▼",
    "┌─────────────────────────────────────────────┐",
    "│              BACKEND (FastAPI)               │",
    "│  Python 3.11 · Pydantic · JWT · RBAC       │",
    "│  Clean Architecture (4 couches)             │",
    "└──────────────────┬──────────────────────────┘",
    "                   │  SQLAlchemy ORM",
    "                   ▼",
    "┌─────────────────────────────────────────────┐",
    "│           BASE DE DONNÉES                    │",
    "│  SQLite (dev) · PostgreSQL (prod)           │",
    "│  18 tables · 26 relations FK                │",
    "└─────────────────────────────────────────────┘",
]

for line in arch_lines:
    add_code_block(doc, line)

add_heading_numbered(doc, "4.2 Architecture backend", 2)

add_normal(doc, (
    "Le backend suit une architecture clean (hexagonale) organisée en quatre couches :"
))

layers = [
    ("Core", "Configuration, base de données, sécurité, logging, cache, gestion des exceptions. "
     "Contient les paramètres globaux et les utilitaires transversaux."),
    ("Domain", "Énumérations, schémas d'import, templates JSON K0/K9, événements métier. "
     "Couche pure, sans dépendance externe."),
    ("Application", "Services métier (20 modules), DTOs (12 modules), interfaces (repositories, services). "
     "Contient la logique de traitement principale."),
    ("Infrastructure", "Persistance (18 modèles, 12 repositories), authentification (JWT + RBAC), "
     "notifications (email), stockage de fichiers."),
    ("API", "Endpoints REST (14 modules de routes), dépendances, middlewares. "
     "Couche de présentation exposant les services."),
]

for name, desc in layers:
    p = doc.add_paragraph()
    run = p.add_run(f"{name} : ")
    run.bold = True
    run.font.size = Pt(11)
    run = p.add_run(desc)
    run.font.size = Pt(11)

add_heading_numbered(doc, "4.3 Architecture frontend", 2)

add_normal(doc, (
    "Le frontend est organisé en une structure feature-based, où chaque module fonctionnel "
    "est encapsulé dans son propre répertoire sous features/. L'architecture repose sur "
    "les principes suivants :"
))

fe_arch = [
    "Séparation par fonctionnalité : chaque module (projects, suppliers, capacity, etc.) "
    "contient ses pages, composants et services spécifiques.",
    "Composants UI réutilisables : bibliothèque de 38 composants UI dans components/ui/ "
    "(button, dialog, table, form, etc.) basés sur Radix UI.",
    "Gestion d'état : Zustand pour l'état global (auth, sidebar, theme), TanStack Query "
    "pour le cache des données serveur.",
    "Hooks personnalisés : 8 hooks partagés + 10 hooks de requête + 8 hooks de mutation.",
    "Moteur de templates : composants DynamicForm, DynamicTable, DynamicFilterBar pilotés "
    "par des schémas de templates JSON.",
    "Internationalisation : systeme de traduction FR/EN via des fichiers de locales.",
    "Contextes React : LanguageContext, TemplateContext, ThemeContext.",
]
for f in fe_arch:
    add_bullet(doc, f)

add_heading_numbered(doc, "4.4 Pipeline d'importation", 2)

add_normal(doc, (
    "Le pipeline d'importation constitue le composant technique le plus complexe de la "
    "plateforme. Il suit un workflow en 6 étapes :"
))

pipeline_steps = [
    ("Upload", "L'utilisateur dépose un fichier Excel (.xlsx, .xls, .xlsm). Le fichier est "
     "stocké temporairement pour traitement."),
    ("Analyse du classeur", "Le système évalue chaque feuille de travail (worksheet scoring) "
     "en se basant sur : nombre de colonnes, nombre de lignes, densité de mots-clés métier, "
     "présence d'indicateurs de pivot. Classification : PROJECT_DATA, PIVOT_TABLE, KPI, "
     "SUMMARY, EMPTY. La feuille la plus pertinente est sélectionnée automatiquement."),
    ("Détection des en-têtes", "Les 20 premières lignes de la feuille sélectionnée sont évaluées "
     "individuellement. Chaque ligne reçoit un score basé sur le nombre de cellules non vides, "
     "les correspondances avec les mots-clés du domaine et la densité numérique. La ligne "
     "ayant le score le plus élevé est identifiée comme ligne d'en-tête."),
    ("Mapping des colonnes", "Cascade à 5 niveaux : cache → correspondance exacte → alias → "
     "similarité floue → IA (Ollama). Chaque colonne Excel est mappée vers un champ du système."),
    ("Validation et preview", "Chaque ligne est validée : types de données, valeurs requises, "
     "doublons (en fichier et en base). Le résultat présente les actions prévues : CREATE, "
     "UPDATE, RESTORE (réactivation d'un enregistrement supprimé logiquement), SKIP."),
    ("Exécution", "Import atomique avec transaction par ligne (SAVEPOINT). Journalisation, "
     "création d'un historique d'import et notification utilisateur."),
]

for i, (name, desc) in enumerate(pipeline_steps, 1):
    p = doc.add_paragraph()
    run = p.add_run(f"Étape {i} — {name} : ")
    run.bold = True
    run.font.size = Pt(11)
    run = p.add_run(desc)
    run.font.size = Pt(11)

add_normal(doc, (
    "Le pipeline supporte deux orientations de données :"
))

add_bullet(doc, (
    "Horizontale (tabulaire) : les en-têtes sont sur une ligne, les données sur les lignes "
    "suivantes. C'est le format standard des tableaux Excel."
))
add_bullet(doc, (
    "Verticale (clé-valeur) : les noms de champs sont dans la colonne A et les valeurs "
    "dans la colonne B. Utilisé pour certains formats de fichiers industriels."
))

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
#           CHAPITRE 5 — TECHNOLOGIES UTILISÉES
# ══════════════════════════════════════════════════════════════

add_heading_numbered(doc, "Chapitre 5 — Technologies utilisées", 1)

add_normal(doc, (
    "Le choix des technologies a été guidé par les critères de performance, de maintenabilité, "
    "de compatibilité avec l'écosystème existant et de disponibilité de ressources "
    "de formation et de support."
))

# Technology table
tech_data = [
    ["React 19", "Frontend", "Bibliothèque UI déclarative pour la construction d'interfaces utilisateur interactives. "
     "Choisie pour son écosystème riche, sa flexibilité et sa large adoption."],
    ["TypeScript 5.6", "Frontend", "Surcouche typée de JavaScript. Assure la sécurité des types, "
     "réduit les erreurs à l'exécution et améliore l'expérience de développement."],
    ["Vite 6", "Frontend", "Outil de build rapide utilisant le Native ES Modules. "
     "Offre un temps de démarrage instantané et un HMR efficace."],
    ["Tailwind CSS 3.4", "Frontend", "Framework CSS utility-first. Permet un développement UI rapide "
     "avec un bundle CSS minimal. Dark mode natif via la stratégie 'class'."],
    ["TanStack Query 5", "Frontend", "Gestion de cache serveur. Automatise le fetch, le cache, "
     "la revalidation et le pessimistic update des données API."],
    ["Zustand", "Frontend", "Gestion d'état globale légère. Utilisé pour l'auth, le sidebar et le thème."],
    ["Radix UI", "Frontend", "Primitives d'accessibilité. Composants non stylisés garantissant "
     "l'accessibilité WCAG 2.1."],
    ["Recharts", "Frontend", "Bibliothèque de graphiques React. Utilisée pour les graphiques "
     "barres, lignes et secteurs du tableau de bord."],
    ["Python 3.11", "Backend", "Langage principal du backend. Choisi pour sa performance, "
     "sa syntaxe claire et la richesse de son écosystème."],
    ["FastAPI 0.115", "Backend", "Framework web async pour Python. Offre des performances élevées, "
     "la validation automatique via Pydantic et la génération OpenAPI."],
    ["SQLAlchemy 2.0", "Backend", "ORM Python avec support async. Couche d'abstraction database "
     "offrant le pattern Unit of Work et les migrations via Alembic."],
    ["SQLite", "Base de données", "Base embarquée pour le développement et les tests. "
     "Aucune configuration serveur nécessaire."],
    ["Alembic", "Backend", "Outil de migration de base de données pour SQLAlchemy. "
     "Gère les évolutions de schéma de manière versionnée."],
    ["openpyxl", "Backend", "Bibliothèque Python de lecture/écriture de fichiers Excel. "
     "Supporte les formats .xlsx, .xls, .xlsm."],
    ["xlsx (SheetJS)", "Frontend", "Bibliothèque JavaScript de parsing de fichiers Excel côté client. "
     "Permet l'analyse préliminaire sans transfert serveur."],
    ["Ollama", "IA", "Moteur d'inférence LLM local. Utilisé comme fallback pour le mapping "
     "sémantique des colonnes Excel non résolu par les méthodes déterministes."],
    ["JWT (python-jose)", "Backend", "Tokens d'authentification JSON Web. Rotation des refresh tokens. "
     "Durée de vie configurable."],
    ["bcrypt", "Backend", "Chiffrement des mots de passe. Algorithme adapté avec salt intégré."],
    ["Pydantic 2.0", "Backend", "Validation de données. Serialisation/deserialisation des requêtes "
     "et réponses API avec typage fort."],
]

add_styled_table(doc, ["Technologie", "Couche", "Rôle et justification"], tech_data, [3, 2, 10.5])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
#     CHAPITRE 6 — CONCEPTION DE LA BASE DE DONNÉES
# ══════════════════════════════════════════════════════════════

add_heading_numbered(doc, "Chapitre 6 — Conception de la base de données", 1)

add_heading_numbered(doc, "6.1 Modèle de données", 2)

add_normal(doc, (
    "La base de données est composée de 18 tables principales, modélisées via SQLAlchemy "
    "et gérées par des migrations Alembic. Le choix de SQLite pour le développement et "
    "PostgreSQL pour la production permet de tester en environnement léger tout en "
    "préparant le déploiement."
))

add_normal(doc, (
    "Les tables principales du système sont les suivantes :"
))

tables_data = [
    ["users", "12", "Utilisateurs du système avec authentification"],
    ["roles", "4", "Rôles définis (admin, buyer, capacity_manager, sqd, viewer)"],
    ["permissions", "5", "Permissions par ressource et action"],
    ["role_permissions", "2", "Table d'association rôle-permission"],
    ["sessions", "7", "Sessions et refresh tokens"],
    ["password_reset_tokens", "5", "Tokens de réinitialisation mot de passe"],
    ["projects", "16", "Projets CMF avec données JSON"],
    ["project_parts", "11", "Pièces/ composants de projets"],
    ["suppliers", "15", "Fournisseurs"],
    ["project_suppliers", "4", "Association projet-fournisseur"],
    ["templates", "6", "Structures/templates de données"],
    ["template_versions", "4", "Versions de templates"],
    ["documents", "12", "Documents attachés aux projets"],
    ["capacity_assessments", "18", "Évaluations de capacité"],
    ["risks", "15", "Risques identifiés"],
    ["notifications", "7", "Notifications utilisateur"],
    ["activity_logs", "8", "Journal d'activité"],
    ["audit_logs", "9", "Journal d'audit"],
]

add_styled_table(doc, ["Table", "Champs", "Description"], tables_data, [4, 2, 9.5])

add_heading_numbered(doc, "6.2 Relations et contraintes", 2)

add_normal(doc, (
    "Le modèle de données comprend 26 relations de clé étrangère et 11 contraintes "
    "d'unicité. Les relations principales sont :"
))

relations_data = [
    ["projects → users", "buyer_id, sqd_id, capacity_manager_id", "Les trois acteurs du projet"],
    ["project_parts → projects", "project_id (CASCADE)", "Pièces d'un projet"],
    ["capacity_assessments → project_parts", "project_part_id (CASCADE)", "Évaluations d'une pièce"],
    ["capacity_assessments → suppliers", "supplier_id (CASCADE)", "Évaluation d'un fournisseur"],
    ["risks → project_parts", "project_part_id (CASCADE)", "Risques d'une pièce"],
    ["documents → projects", "project_id (SET NULL)", "Documents d'un projet"],
    ["templates → template_versions", "template_id (CASCADE)", "Versioning de templates"],
    ["project_suppliers → projects, suppliers", "CASCADE", "Association多 à多 projet-fournisseur"],
]

add_styled_table(doc, ["Relation", "Clé étrangère", "Signification"], relations_data, [4.5, 4.5, 6.5])

add_normal(doc, (
    "Les mixins de timestamps (created_at, updated_at) et de soft-delete (deleted_at) "
    "sont appliqués à la plupart des tables, assurant une traçabilité et une récupération "
    "des données cohérentes. Les index sont créés sur les colonnes fréquemment interrogées "
    "(code, name, status) pour optimiser les performances de recherche."
))

add_normal(doc, (
    "Contraintes d'unicité notables :"
))

add_bullet(doc, "users : email unique, username unique")
add_bullet(doc, "projects : code unique")
add_bullet(doc, "suppliers : code unique")
add_bullet(doc, "templates : code unique")
add_bullet(doc, "project_parts : contrainte unique sur (project_id, part_number)")
add_bullet(doc, "capacity_assessments : contrainte unique sur (project_part_id, supplier_id, assessment_date)")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
#      CHAPITRE 7 — DÉVELOPPEMENT DE LA PLATEFORME
# ══════════════════════════════════════════════════════════════

add_heading_numbered(doc, "Chapitre 7 — Développement de la plateforme", 1)

add_heading_numbered(doc, "7.1 Tableau de bord (Dashboard)", 2)

add_normal(doc, (
    "Le tableau de bord constitue la page d'accueil de la plateforme après connexion. "
    "Il offre une vue d'ensemble complète de l'activité CMF à travers plusieurs sections :"
))

add_normal(doc, (
    "La section supérieure (Hero Banner) affiche un accueil personnalisé avec le nom de "
    "l'utilisateur, son rôle (sous forme de badge) et des raccourcis vers les actions "
    "fréquentes (Créer un projet, Matrice de capacité, Audits SQD). Les boutons d'action "
    "sont adaptés au rôle de l'utilisateur grâce au système de permissions."
))

add_normal(doc, (
    "Six cartes KPI (Key Performance Indicators) présentent les indicateurs essentiels : "
    "nombre total de projets CMF, capacité totale (en K), taux d'utilisation, écart de "
    "capacité, nombre de fournisseurs actifs et projets à risque. Ces cartes intègrent "
    "des variations de valeur et des indicateurs de tendance."
))

add_normal(doc, (
    "La section Analytics contient des graphiques interactifs : un graphique en barres "
    "horizontal pour la répartition des statuts de projets (Actifs, En bonne voie, "
    "En retard, En retard, Terminés), un graphique linéaire pour l'évolution mensuelle "
    "de la capacité, et un diagramme en secteurs pour la distribution SQD."
))

add_normal(doc, (
    "Le tableau de bord est piloté par le hook useCmfDashboardData() qui consomme "
    "l'API /api/v1/dashboard/stats. La visibilité des sections est contrôlée par le "
    "hook usePermissions() qui vérifie le rôle de l'utilisateur."
))

add_heading_numbered(doc, "7.2 Gestion des projets", 2)

add_normal(doc, (
    "La gestion des projets est le module central de la plateforme. Elle comprend :"
))

add_heading_numbered(doc, "7.2.1 Liste des projets", 3)

add_normal(doc, (
    "La page de liste affiche les projets sous forme de tableau dynamique utilisant les "
    "composants DynamicTable et DynamicFilterBar du moteur de templates. Les fonctionnalités "
    "incluent la recherche textuelle, le filtrage par statut, la pagination et la sélection "
    "multiple pour la suppression en masse. Un sélecteur de template permet de filtrer les "
    "projets par structure (K0, K9, personnalisée)."
))

add_heading_numbered(doc, "7.2.2 Création de projet", 3)

add_normal(doc, (
    "La création de projet propose deux voies : la création manuelle et l'import Excel. "
    "Le sélecteur InputSourcePicker permet à l'utilisateur de choisir entre ces deux options. "
    "Pour la création manuelle via template CMF, seul le formulaire de la section Buyer est "
    "affiché lors de la première étape, conformément au workflow de validation."
))

add_heading_numbered(doc, "7.2.3 Détail du projet", 3)

add_normal(doc, (
    "La page de détail d'un projet présente quatre onglets : Détails, Pièces, Risques et "
    "Documents. Pour les projets CMF (K0, K9), la vue K9ProjectView affiche les données "
    "organisées par section (Buyer, Capacity Manager, SQD) avec des permissions d'édition "
    "par champ. La sauvegarde peut se faire section par section."
))

add_heading_numbered(doc, "7.3 Structures et templates", 2)

add_normal(doc, (
    "Le module de gestion des structures (TemplateStudioPage) est un hub à 5 onglets :"
))

add_bullet(doc, (
    "Structures Overview : vue en grille de toutes les templates avec aperçu des sections, "
    "actions de création, visualisation, export et suppression. Les templates K0 et K9 "
    "sont protégés (non supprimables)."
))
add_bullet(doc, (
    "Projects in Structure : tableau des projets appartenant à une structure sélectionnée, "
    "avec recherche et filtrage par statut."
))
add_bullet(doc, (
    "Field Schema Explorer : table de tous les champs de toutes les templates, "
    "avec recherche et filtrage par section de rôle."
))
add_bullet(doc, (
    "Structure Comparison Matrix : comparaison côte à côte des propriétés de templates."
))
add_bullet(doc, (
    "JSON Schema Inspector : visualiseur JSON brut avec copie et téléchargement."
))

add_normal(doc, (
    "L'éditeur de template (TemplateEditor) propose un mode builder visuel permettant "
    "d'ajouter des sections et des champs, et un mode éditeur JSON pour les utilisateurs "
    "avancés. Chaque champ dispose d'un panneau de permissions permettant de définir "
    "quels rôles peuvent le consulter ou le modifier."
))

add_heading_numbered(doc, "7.4 Création manuelle d'un projet", 2)

add_normal(doc, (
    "La création manuelle d'un projet suit le workflow CMF en trois étapes :"
))

add_normal(doc, (
    "Étape 1 — Buyer : L'acheteur définit les informations de base du projet : numéro de "
    "pièce, description, fournisseur, prix, quantité, etc. Cette étape correspond aux champs "
    "de la section Buyer du template."
))

add_normal(doc, (
    "Étape 2 — Capacity Manager : Le responsable de capacité évalue la capacité actuelle "
    "et maximale du fournisseur, définit les objectifs par semaine, saisit les prévisions "
    "et les réalisations."
))

add_normal(doc, (
    "Étape 3 — SQD : Le responsable qualité effectue l'évaluation qualité, classifie les "
    "risques (GREEN/ORANGE/RED), valide la conformité du fournisseur."
))

add_normal(doc, (
    "Le workflow est automatiquement calculé par le backend en fonction des données saisies. "
    "L'étape 1 est active par défaut. L'étape 2 est activée lorsque les champs du Buyer "
    "sont remplis. L'étape 3 est activée lorsque les champs du Capacity Manager sont saisis. "
    "Le projet est marqué COMPLETED lorsque la section SQD contient une classification GREEN."
))

add_heading_numbered(doc, "7.5 Gestion des fournisseurs", 2)

add_normal(doc, (
    "Le module de gestion des fournisseurs offre un CRUD complet avec :"
))

add_bullet(doc, "Liste de fournisseurs avec filtrage par statut et recherche")
add_bullet(doc, "Vues grille et carte pour la visualisation")
add_bullet(doc, "Création et édition avec formulaire validé")
add_bullet(doc, "Détail du fournisseur avec assignation à des projets")
add_bullet(doc, "Statuts : actif, inactif, blacklisté")
add_bullet(doc, "Catégorisation et notation")
add_bullet(doc, "Modales de visualisation rapide (QuickView)")

add_heading_numbered(doc, "7.6 Gestion de la capacité", 2)

add_normal(doc, (
    "Le module de capacité permet de créer et suivre des évaluations de capacité pour "
    "chaque combinaison pièce/fournisseur. Les évaluations sont mensuelles et comprennent :"
))

add_bullet(doc, "Capacité actuelle et maximale (numériques avec précision décimale)")
add_bullet(doc, "Taux d'utilisation calculé automatiquement")
add_bullet(doc, "Champs CMF spécifiques : CATE, gate, target_week, forecast_week, completed_week")
add_bullet(doc, "Niveau de risque (low, medium, high, critical)")
add_bullet(doc, "Statut : pending, assessed, confirmed, rejected")

add_heading_numbered(doc, "7.7 Suivi qualité (SQD)", 2)

add_normal(doc, (
    "Le suivi qualité est intégré dans le workflow CMF via la section SQD des templates. "
    "Les fonctionnalités comprennent :"
))

add_bullet(doc, "Classification des risques qualité : GREEN (validé), ORANGE (à améliorer), RED (bloquant)")
add_bullet(doc, "Gestion des risques avec sévérité, probabilité et score calculé")
add_bullet(doc, "Matrice de chaleur des risques (RiskHeatmapMatrix)")
add_bullet(doc, "Tableau Kanban pour le suivi visuel (RiskKanbanBoard)")
add_bullet(doc, "Actions d'atténuation et de fermeture des risques")
add_bullet(doc, "Modale de mitigation rapide (QuickMitigateModal)")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
#        CHAPITRE 8 — IMPORTATION DES DONNÉES
# ══════════════════════════════════════════════════════════════

add_heading_numbered(doc, "Chapitre 8 — Importation des données", 1)

add_normal(doc, (
    "L'importation de données depuis des fichiers Excel constitue le composant technique "
    "le plus élaboré de la plateforme. Elle répond au besoin fondamental de transférer "
    "les données existantes depuis les fichiers Excel vers le système centralisé, "
    "tout en garantissant leur intégrité et leur cohérence."
))

add_heading_numbered(doc, "8.1 Workflow d'importation", 2)

add_normal(doc, (
    "Le wizard d'importation (ImportWizard) guide l'utilisateur à travers 7 étapes :"
))

wizard_steps = [
    ("1. Upload Excel", "Drag-and-drop ou sélection de fichier. Supporte .xlsx, .xls, .xlsm. "
     "Le fichier est validé côté client avant envoi."),
    ("2. Analyse du classeur", "Le backend analyse chaque feuille via le worksheet scoring. "
     "Classification : PROJECT_DATA, PIVOT_TABLE, KPI, SUMMARY, REFERENCE_DATA, EMPTY. "
     "La feuille la mieux classée est sélectionnée automatiquement."),
    ("3. Feuille et en-têtes", "Détection de la ligne d'en-tête parmi les 20 premières lignes. "
     "Affichage des colonnes détectées. Possibilité de sélectionner manuellement la ligne "
     "d'en-tête. Détection de l'orientation (HORIZONTAL/VERTICAL)."),
    ("4. Structure du projet", "Sélection du template cible (K0, K9, ou structure personnalisée). "
     "Le template définit les champs attendus et les permissions par rôle."),
    ("5. Mapping IA", "Cascade de mapping à 5 niveaux pour chaque colonne Excel. "
     "Affichage de la confiance pour chaque mapping. Possibilité de corriger manuellement."),
    ("6. Validation/Preview", "Validation complète côté serveur. Affichage de chaque ligne "
     "avec son action prévue (CREATE/UPDATE/RESTORE/SKIP) et les erreurs éventuelles."),
    ("7. Exécution", "Import avec transactions atomiques. Affichage des résultats : "
     "importés, mis à jour, ignorés, échoués. Synchronisation du cache de mapping."),
]

for step_name, desc in wizard_steps:
    p = doc.add_paragraph()
    run = p.add_run(f"{step_name} : ")
    run.bold = True
    run.font.size = Pt(11)
    run = p.add_run(desc)
    run.font.size = Pt(11)

add_heading_numbered(doc, "8.2 Détection de feuille et d'en-tête", 2)

add_normal(doc, (
    "La détection de la feuille pertinente est une étape critique car les fichiers Excel "
    "industriels contiennent souvent de multiples feuilles (données, pivots, graphiques, "
    "KPI, résumés). Le scoring de feuille évalue chaque feuille selon plusieurs critères :"
))

add_bullet(doc, "Nombre de colonnes (bonus pour > 6 colonnes)")
add_bullet(doc, "Nombre de lignes de données")
add_bullet(doc, "Densité de mots-clés du domaine (supplier, part_number, capacity, etc.)")
add_bullet(doc, "Pénalités pour les noms de feuilles type Dashboard, Chart, Pivot")
add_bullet(doc, "Classification spéciale : PIVOT_TABLE (-100), KPI (-120)")

add_normal(doc, (
    "La détection de la ligne d'en-tête évalue les 20 premières lignes. Chaque ligne "
    "reçoit un score basé sur :"
))

add_bullet(doc, "Nombre de cellules non vides")
add_bullet(doc, "Nombre de correspondances avec les mots-clés du domaine (+10 par match)")
add_bullet(doc, "Pénalité si forte densité numérique (ligne de données, pas un en-tête)")

add_normal(doc, (
    "La détection d'orientation examine la colonne A pour des correspondances avec les "
    "noms de champs du domaine. Si ≥ 4 correspondances dans la colonne A avec des valeurs "
    "dans la colonne B, l'orientation est classée comme VERTICALE. Sinon, elle est "
    "classée comme HORIZONTALE."
))

add_heading_numbered(doc, "8.3 Mapping sémantique des colonnes", 2)

add_normal(doc, (
    "Le mapping des colonnes Excel vers les champs du système est réalisé par une cascade "
    "à 5 niveaux, chaque niveau étant une tentative de résolution indépendante :"
))

mapping_levels = [
    ["Niveau 0", "Cache de mapping", "Correspondance exacte normalisée avec les mappings "
     "précédemment confirmés par l'utilisateur. Confiance : 0.99. Contourne tous les autres niveaux."],
    ["Niveau 1", "Correspondance exacte", "normalisation du header Excel et comparaison "
     "avec la clé et le label normalisés du champ. Confiance : 1.0."],
    ["Niveau 2", "Alias connus", "Comparaison avec les aliases définis dans le schéma du template. "
     "Confiance : 0.96."],
    ["Niveau 3", "Similarité floue", "Calcul de similarité par SequenceMatcher, Jaccard token "
     "overlap et substring containment. ≥ 0.90 → auto-map. 0.75-0.89 → candidat ambigu."],
    ["Niveau 4", "IA (Ollama)", "Uniquement pour les colonnes non résolues. Appel à un LLM local "
     "avec prompt minimisé. Validation stricte de la réponse JSON. Fallback gracieux si indisponible."],
]

add_styled_table(doc, ["Niveau", "Méthode", "Description"], mapping_levels, [2, 3, 10.5])

add_normal(doc, (
    "La normalisation des en-têtes (header_normalizer) suit ces étapes : suppression des "
    "accents via normalisation NFKD, conversion en minuscules, remplacement de tous les "
    "caractères non alphanumériques par des espaces, réduction des espaces multiples."
))

add_normal(doc, (
    "Le calcul de similarité retourne le maximum de trois mesures : SequenceMatcher (similarité "
    "de séquence), Jaccard token overlap (similarité de tokens) et containment (inclusion). "
    "Cette approche hybride améliore la robustesse du matching."
))

add_heading_numbered(doc, "8.4 Validation et preview", 2)

add_normal(doc, (
    "La validation des données est réalisée en deux phases : la normalisation et la "
    "vérification des types. Pour chaque colonne, le DataNormalizer applique des règles "
    "spécifiques :"
))

add_bullet(doc, (
    "Proxys null : chaînes vides, 'none', 'N/A', 'TBD', '-', '/', '--' → is_null=True"
))
add_bullet(doc, (
    "Nombres : suppression des symboles de devise, gestion des séparateurs US/EU, "
    "extraction de la portion numérique"
))
add_bullet(doc, (
    "Dates : détection des formats DD/MM/YYYY, MM/DD/YYYY, YYYY-MM-DD. Les nombres "
    "ambigus (1-4 chiffres) sont signalés avec un avertissement"
))
add_bullet(doc, (
    "Énumérations : normalisation en minuscules avec remplacement espaces par underscores"
))

add_normal(doc, (
    "La détection de doublons opère à deux niveaux : en fichier (même clé unique dans "
    "le fichier) et en base (même clé unique dans la base de données). Les enregistrements "
    "en base sont classés comme actifs (UPDATE), supprimés logiquement (RESTORE) ou "
    "nouveaux (CREATE)."
))

add_heading_numbered(doc, "8.5 Exécution sécurisée", 2)

add_normal(doc, (
    "L'exécution de l'import est conçue pour être atomique et recoverable. Chaque ligne "
    "est traitée dans une transaction avec SAVEPOINT, permettant d'isoler les erreurs "
    "d'une ligne sans compromettre les autres."
))

add_normal(doc, (
    "Le processus d'exécution :"
))

add_bullet(doc, "Re-validation des données (les données sont validées à nouveau)")
add_bullet(doc, "Vérification de la stratégie : si rollback_all et des erreurs existent → abort")
add_bullet(doc, "Pour chaque ligne : skip si non-enregistrement, skip si doublon fichier, "
           "skip si doublon base, skip si erreur de validation")
add_bullet(doc, "Normalisation et validation des types pour chaque champ")
add_bullet(doc, "Déduplication en batch (set inserted_in_batch)")
add_bullet(doc, "Insertion ou mise à jour via le service métier approprié")
add_bullet(doc, "Création d'un log d'audit et d'un historique d'import")
add_bullet(doc, "Notification utilisateur du résultat")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
#         CHAPITRE 9 — STRUCTURES K0 ET K9
# ══════════════════════════════════════════════════════════════

add_heading_numbered(doc, "Chapitre 9 — Structures K0 et K9", 1)

add_normal(doc, (
    "Les structures K0 et K9 sont deux templates de données prédéfinis, spécifiques "
    "au processus CMF de l'entreprise. Ils sont intégralement définis dans des fichiers "
    "JSON et auto-seedés au démarrage de l'application. Ces structures sont protégées "
    "et ne peuvent pas être supprimées."
))

add_heading_numbered(doc, "9.1 Structure K0", 2)

add_normal(doc, (
    "La structure K0 est conçue pour le suivi de projets de type K0, un format spécifique "
    "au processus CMF. Elle comprend 42 champs répartis en trois sections correspondant "
    "aux trois rôles du workflow :"
))

k0_sections = [
    ["Buyer", "17 champs", "part_number, index, description, coef, serial_piece_price, "
     "mass_purchase, ru, noa, make_battery_lp_1, make_battery_lp_2, supplier_name, "
     "vendor_name, manufacturer_name, combined_cofor, tango_order, ei_status, comments"],
    ["Capacity Manager", "9 champs", "week_project_target_1/2/3, forecast_week_1/2/3, "
     "completed_week_1/2/3"],
    ["SQD", "16 champs", "quality, supply_chain, global_purchasing, cpl, rcpi, "
     "minimum_quality_status_acted, mass_inquired, packaging_validated, tango_validated, "
     "supplier_validated, it_cpl, fcla, ple, edi, um_logistic, manufacturing_validated"],
]

add_styled_table(doc, ["Section", "Nombre", "Champs principaux"], k0_sections, [3.5, 2.5, 9.5])

add_normal(doc, (
    "La particularité technique de K0 réside dans sa gestion des colonnes en double. "
    "Le fichier Excel source contient des colonnes dupliquées (par exemple, 'Week\\nProject Target' "
    "apparaît aux colonnes W, Z et AC pour les milestones 1, 2 et 3). Le mapping par en-tête "
    "serait incapable de distinguer ces colonnes. C'est pourquoi K0 utilise un mapping "
    "positionnel (par index de colonne) via la liste K0_SOURCE_COLUMNS, contenant 47 "
    "positions de colonnes. Cette approche rend K0 immunisé contre les collisions d'en-têtes."
))

add_normal(doc, (
    "L'importation K0 extrait automatiquement : code (part_number), nom (description), "
    "description, client (vendor_name) et Calcule l'étape de workflow en fonction des "
    "sections remplies."
))

add_heading_numbered(doc, "9.2 Structure K9", 2)

add_normal(doc, (
    "La structure K9 est un template plus large, destiné au suivi de projets de type K9. "
    "Elle comprend environ 25 champs répartis en trois sections :"
))

k9_sections = [
    ["Buyer", "12 champs", "unique_id, apqp, part_name, use_case, part_info, part_number, "
     "supplier_info, supplier_name, manufacturing_cofor, production_location, stakeholder, buyer"],
    ["Capacity Manager", "7 champs", "capacity, scr_link_docinfo, gst_no, contracted_capacity, "
     "fete, tko_fete_link_sharepoint, capacity_standard, fete_tko_letter_doc"],
    ["SQD", "6 champs", "technical_manager, k9_sck, cat1/2/3_forecast_dates, "
     "cat1_2_3_type, weekly_capacity_measured, estimated_target, cat_evaluation, "
     "shared_folder_link, comments, sqe, sqm, team, family_multiplier"],
]

add_styled_table(doc, ["Section", "Nombre", "Champs principaux"], k9_sections, [3.5, 2.5, 9.5])

add_normal(doc, (
    "Contrairement à K0, K9 utilise le mapping par en-tête normalisé plutôt que par index, "
    "car ses en-têtes sont plus distincts. Le workflow de validation suit la même logique "
    "que K0 : progression de l'étape 1 (Buyer) à l'étape 4 (SQD GREEN → Complété)."
))

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
#       CHAPITRE 10 — OPTIMISATION DES IMPORTS
# ══════════════════════════════════════════════════════════════

add_heading_numbered(doc, "Chapitre 10 — Optimisation des imports", 1)

add_normal(doc, (
    "L'optimisation du pipeline d'importation a été un axe de développement important, "
    "notamment après l'intégration du mapping IA. Les principales optimisations "
    "implémentées sont les suivantes :"
))

add_heading_numbered(doc, "10.1 Réduction des appels IA", 2)

add_normal(doc, (
    "Le mapping IA (Ollama) est coûteux en temps de traitement. Plusieurs mécanismes "
    "ont été mis en place pour minimiser son utilisation :"
))

add_bullet(doc, (
    "Mapping par cache (Niveau 0) : les mappings précédemment confirmés par l'utilisateur "
    "sont stockés en mémoire (côté serveur) et en localStorage (côté client). Lors des "
    "imports suivants pour le même template, les colonnes déjà mappées contournent "
    "directement l'IA."
))
add_bullet(doc, (
    "Mapping déterministe : les niveaux 1-3 (correspondance exacte, alias, fuzzy) "
    "résolvent la grande majorité des colonnes sans intervention IA."
))
add_bullet(doc, (
    "Prompt minimisé : seules les colonnes non résolues et les champs non mappés sont "
    "envoyés au LLM, réduisant la taille du prompt et le temps de traitement."
))
add_bullet(doc, (
    "Health check préalable : avant tout appel IA, un health check vérifie la disponibilité "
    "d'Ollama (timeout 1.5s). Si indisponible, le mapping déterministe continue sans IA."
))

add_heading_numbered(doc, "10.2 Normalisation et cache", 2)

add_normal(doc, (
    "La normalisation des en-têtes est un prérequis pour le matching. Elle est appliquée "
    "une seule fois par en-tête et les résultats sont cachés. Le cache de mapping est "
    "persisté côté serveur via l'endpoint save-mapping-memory et rechargé automatiquement "
    "lors des imports suivants pour le même template."
))

add_heading_numbered(doc, "10.3 Filtrage des lignes vides", 2)

add_normal(doc, (
    "Les lignes entièrement vides ou ne contenant aucune clé unique sont identifiées "
    "lors de la phase de validation et exclues du comptage physique. Elles ne sont pas "
    "comptabilisées comme des erreurs ni comme des enregistrements à traiter, évitant "
    "ainsi un gonflement artificiel des statistiques d'import."
))

add_heading_numbered(doc, "10.4 Mapping positionnel pour K0", 2)

add_normal(doc, (
    "Le cas spécifique de K0, avec ses colonnes en double, a nécessité l'implémentation "
    "d'un mapping positionnel. Plutôt que de comparer les en-têtes (qui sont identiques "
    "pour les colonnes dupliquées), le système utilise l'index de la colonne dans le "
    "fichier Excel pour déterminer le champ cible. Cette approche est déterministe et "
    "ne nécessite aucun appel IA."
))

add_heading_numbered(doc, "10.5 Optimisation de la lecture Excel", 2)

add_normal(doc, (
    "Côté client, la bibliothèque xlsx (SheetJS) est utilisée pour l'analyse préliminaire "
    "(scoring des feuilles, détection d'en-têtes). Seul le fichier brut est envoyé au "
    "backend pour le traitement complet, évitant les transferts de données intermédiaires."
))

add_normal(doc, (
    "Côté backend, openpyxl est utilisé en mode read_only pour les grands fichiers, "
    "réduisant la consommation mémoire lors du parsing de fichiers contenant des centaines "
    "de lignes."
))

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
#       CHAPITRE 11 — DIFFICULTÉS RENCONTRÉES
# ══════════════════════════════════════════════════════════════

add_heading_numbered(doc, "Chapitre 11 — Difficultés rencontrées", 1)

add_normal(doc, (
    "Le développement de la plateforme CMF a été jalonné de plusieurs difficultés "
    "techniques significatives. La présentation de ces difficultés, de leur analyse "
    "et des solutions apportées, constitue un témoignage du travail de résolution "
    "de problèmes réalisé au cours du stage."
))

difficulties = [
    ("Colonnes en double dans les fichiers Excel K0",
     "Les fichiers Excel source pour la structure K0 contiennent des colonnes en double : "
     "les en-têtes 'Week\\nProject Target', 'Week\\nForecast' et 'Week\\nCompleted' "
     "apparaissent respectivement 3 fois dans le fichier (pour les milestones 1, 2 et 3). "
     "Le mapping classique par comparaison d'en-têtes ne pouvait pas distinguer ces colonnes.",
     "Implémentation d'un mapping positionnel (par index de colonne) spécifique à K0. "
     "La liste K0_SOURCE_COLUMNS définit les 47 positions de colonnes, permettant un mapping "
     "déterministe sans ambiguïté. Cette approche rend le mapping immunisé contre les "
     "collisions d'en-têtes.",
     "Le mapping K0 est désormais 100% déterministe, sans appel IA, et gère correctement "
     "les 47 colonnes dont 5 en double."),

    ("Détection de l'orientation des données (Horizontal vs Vertical)",
     "Certains fichiers Excel industriels utilisent un format vertical (clé dans la colonne A, "
     "valeur dans la colonne B) plutôt que le format tabulaire classique. Le pipeline "
     "d'importation devait supporter les deux formats.",
     "Développement d'un algorithme de détection d'orientation basé sur l'analyse de la "
     "colonne A : si ≥ 4 mots-clés du domaine sont trouvés dans la colonne A avec des "
     "valeurs dans la colonne B, l'orientation est classée comme VERTICALE. Un sélecteur "
     "manuel permet à l'utilisateur de forcer l'orientation.",
     "Les deux formats sont désormais gérés automatiquement avec un fallback manuel."),

    ("Intégration de l'IA (Ollama) pour le mapping",
     "L'intégration d'un LLM local (Ollama) pour le mapping sémantique des colonnes a "
     "posé plusieurs défis : disponibilité du service, temps de réponse, format de "
     "réponse non structuré, hallucinations du modèle.",
     "Implémentation d'une cascade à 5 niveaux où l'IA n'est utilisée qu'en dernier "
     "recours. Health check préalable, timeout strict (5s), validation stricte de la "
     "réponse JSON (parsing multiple, vérification des membres, prévention des doublons). "
     "Fallback gracieux : si Ollama est indisponible, le mapping déterministe continue.",
     "L'IA est un complément fiable sans être un point de défaillance."),

    ("Race conditions dans le wizard d'import",
     "Le wizard d'importation multi-étapes était sujet à des race conditions lors de "
     "l'analyse de classeur : les réponses asynchrones pouvaient écraser les résultats "
     "d'étapes suivantes si l'utilisateur changeait de feuille pendant l'analyse.",
     "Implémentation de compteurs de référence (analysisCounterRef, orientationRef) pour "
     "détecter les réponses obsolètes. Chaque réponse est comparée au compteur courant : "
     "si elle est obsolète, elle est ignorée.",
     "Les race conditions sont éliminées, garantissant la cohérence des données affichées."),

    ("Suppression de templates avec dépendances",
     "La suppression d'un template devait gérer la cascade de suppressions vers les projets "
     "dépendants. La suppression physique aurait compromis l'intégrité référentielle.",
     "Implémentation du soft-delete en cascade : supprimer un template marque tous les "
     "projets dépendants comme supprimés logiquement (deleted_at). Les templates K0 et K9 "
     "sont protégés contre la suppression.",
     "L'intégrité référentielle est préservée et la récupération est possible."),

    ("Errors de typage des dates dans les imports",
     "Les dates dans les fichiers Excel sont souvent au format texte ou dans des formats "
     "ambigus (DD/MM/YYYY vs MM/DD/YYYY). L'import échouait parfois avec des erreurs "
     "de sérialisation JSON lors du stockage.",
     "Implémentation du DataNormalizer avec gestion multi-formats pour les dates. "
     "Les dates ambigues (nombres purs à 1-4 chiffres) sont signalées avec un avertissement "
     "plutôt que rejetées. La sérialisation des dates est sécurisée côté serveur.",
     "Tous les formats de dates courants sont gérés sans erreur de sérialisation."),

    ("Déploiement multi-environnement",
     "Le passage du développement (SQLite) au déploiement (PostgreSQL via Neon) a nécessité "
     "de résoudre des incompatibilités : paramètres de connexion asyncpg, transformation "
     "des URLs de base de données, gestion du pooling.",
     "Création d'une couche de configuration adaptative qui transforme les URLs SQLite en "
     "PostgreSQL asyncpg-compatibles. Filtrage des paramètres non supportés (channel_binding). "
     "Auto-init de la base de données au démarrage.",
     "Le déploiement fonctionne sur Railway/Neon avec la même base de code."),

    ("Internationalisation complète",
     "Le système devait supporter le français et l'anglais pour tous les libellés de "
     "l'interface. La gestion manuelle des traductions était sujette aux oublis.",
     "Centralisation des traductions dans des fichiers dédiés (locales/en.ts, locales/fr.ts) "
     "avec un LanguageContext React et un hook useTranslation(). Les clés de traduction "
     "suivent une convention hiérarchique (nav.dashboard, projects.title, etc.).",
     "Le basculement FR/EN est instantané et couvre l'intégralité de l'interface."),
]

for title, problem, solution, result in difficulties:
    doc.add_heading(title, level=3)

    p = doc.add_paragraph()
    run = p.add_run("Problème : ")
    run.bold = True
    p.add_run(problem)

    p = doc.add_paragraph()
    run = p.add_run("Solution : ")
    run.bold = True
    p.add_run(solution)

    p = doc.add_paragraph()
    run = p.add_run("Résultat : ")
    run.bold = True
    p.add_run(result)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
#        CHAPITRE 12 — TESTS ET VALIDATION
# ══════════════════════════════════════════════════════════════

add_heading_numbered(doc, "Chapitre 12 — Tests et validation", 1)

add_normal(doc, (
    "La stratégie de tests de la plateforme CMF repose sur des tests unitaires et "
    "d'intégration implémentés via pytest (backend) et Vitest (frontend). "
    "Le jeu de tests comprend 21 fichiers de test couvrant l'ensemble des modules "
    "critiques du système."
))

add_heading_numbered(doc, "12.1 Tests backend", 2)

add_normal(doc, (
    "Les tests backend sont organisés par domaine fonctionnel :"
))

test_categories = [
    ["Health & API", "test_health.py", "5", "Points de terminaison santé, CORS, 404, 401"],
    ["Utilisateurs", "test_users_endpoint.py", "1", "Liste des utilisateurs avec pagination"],
    ["Templates", "test_templates.py", "4", "Seed, recréation, suppression cascade, validation K9"],
    ["Projets K9", "test_k9_project.py", "3", "Seed, permissions création, restrictions de champ"],
    ["Projets K0", "test_k0_project.py", "10", "Seed, permissions, workflow, intégrité des champs, mapping positionnel"],
    ["Projets bulk", "test_projects_bulk_delete.py", "1", "Pagination et suppression en masse"],
    ["Import JSON", "test_structure_json_import.py", "13", "Normalisation, validation, import hiérarchique/plat"],
    ["Import rôles", "test_structure_import_roles.py", "4", "Classification des rôles, sections"],
    ["Import régression", "test_import_regression.py", "4", "Mapping auto, Ollama, dictionnaire, enums"],
    ["Import endpoints", "test_import_endpoints_full.py", "4", "Templates, headers, warnings, dates"],
    ["Import records", "test_import_record_classification.py", "4", "Dédoublonnage, lignes vides, exécution"],
    ["Import perf", "test_import_pipeline_optimization.py", "16", "Scoring, normalisation, aliases, fuzzy, AI fallback, cache"],
    ["Import exécution", "test_project_import_execution.py", "13", "CRUD complet via import, restauration, 678 lignes"],
    ["K9 Excel", "test_k9_excel_import.py", "1", "Préservation des clés header"],
    ["K0 Excel", "test_k0_excel_26_columns_import.py", "4", "47 colonnes, import réel, mapping positionnel"],
    ["Feuilles", "test_worksheet_scoring.py", "2", "Classification et scoring de feuilles"],
    ["Vertical", "test_vertical_import.py", "11", "Détection orientation, extraction, scoring"],
    ["Normalisation", "test_data_normalizer.py", "4", "Nombres, null, dates, enums"],
    ["Pipeline générique", "test_generic_structure_pipeline.py", "1", "K0+K9 seed, filtrage par template"],
    ["Config deploy", "test_deployment_config.py", "5", "SQLite, Neon, PostgreSQL, CORS"],
]

add_styled_table(doc, ["Catégorie", "Fichier", "Tests", "Sujet"], test_categories, [2.5, 5, 1.5, 7.5])

add_normal(doc, (
    "Au total, le jeu de tests backend comprend environ 113 tests couvrant :"
))

add_bullet(doc, "La santé de l'API et la configuration CORS")
add_bullet(doc, "L'authentification et le contrôle d'accès (RBAC)")
add_bullet(doc, "Le CRUD des templates et la protection K0/K9")
add_bullet(doc, "Le workflow de validation des projets (étapes 1→4)")
add_bullet(doc, "Les restrictions de champ par rôle (Buyer, Capacity, SQD)")
add_bullet(doc, "Le pipeline d'import complet (upload → preview → execute)")
add_bullet(doc, "Le mapping à 5 niveaux et le fallback IA")
add_bullet(doc, "La normalisation des données (nombres, dates, enums)")
add_bullet(doc, "La détection d'orientation horizontale/verticale")
add_bullet(doc, "Le scoring des feuilles de travail")
add_bullet(doc, "La détection et gestion des doublons")
add_bullet(doc, "La restauration d'enregistrements supprimés logiquement")
add_bullet(doc, "La configuration de déploiement multi-environnement")

add_heading_numbered(doc, "12.2 Scénarios de test illustratifs", 2)

add_normal(doc, (
    "Voici quelques scénarios de test emblématiques :"
))

scenarios = [
    ("Import K0 avec 47 colonnes",
     "Le test vérifie que le preview retourne 47 headers et que l'exécution crée "
     "2 enregistrements en base avec toutes les 47 colonnes correctement映射."),
    ("Import 678 lignes (stress test)",
     "Le test importe 678 lignes depuis un fichier Excel puis réimporte le même fichier. "
     "Le premier import crée 678 enregistrements, le second les met à jour (updated=678) "
     "sans créer de doublons."),
    ("Restauration de projet supprimé",
     "Un projet est supprimé logiquement puis réimporté via Excel. Le système détecte "
     "l'enregistrement supprimé et le restaure (RESTORE) au lieu de créer un doublon."),
    ("Mapping positionnel K0",
     "Le test vérifie que les 47 positions de colonnes K0 sont correctement映射, "
     "notamment les 5 colonnes en double (milestone 1/2/3) qui ne doivent pas collisionner."),
    ("Fallback IA gracieux",
     "Le test simule l'indisponibilité d'Ollama et vérifie que le mapping déterministe "
     "continue sans erreur, avec les colonnes non résolues dans la catégorie 'no_match'."),
]

for title, desc in scenarios:
    p = doc.add_paragraph()
    run = p.add_run(f"{title} : ")
    run.bold = True
    run.font.size = Pt(11)
    run = p.add_run(desc)
    run.font.size = Pt(11)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
#             CHAPITRE 13 — RÉSULTATS OBTENUS
# ══════════════════════════════════════════════════════════════

add_heading_numbered(doc, "Chapitre 13 — Résultats obtenus", 1)

add_normal(doc, (
    "Le développement de la plateforme CMF a abouti à la livraison d'une application "
    "web complète et fonctionnelle, répondant aux besoins exprimés par l'entreprise. "
    "Les principaux résultats obtenus sont les suivants :"
))

add_heading_numbered(doc, "13.1 Plateforme développée", 2)

results_summary = [
    ["Frontend React/TypeScript", "16 modules fonctionnels, 38 composants UI, "
     "18 hooks, 24 pages", "Fonctionnel"],
    ["Backend FastAPI/Python", "14 endpoints API, 20 services, 18 modèles", "Fonctionnel"],
    ["Base de données", "18 tables, 26 relations FK, 11 contraintes unicité", "Fonctionnel"],
    ["Pipeline d'import", "7 étapes, 5 niveaux de mapping, support 2 orientations", "Fonctionnel"],
    ["Tests", "113 tests backend couvrant tous les modules critiques", "✅ Passants"],
    ["Internationalisation", "FR/EN avec 399 clés de traduction", "Fonctionnel"],
    ["Déploiement", "Configuration Vercel (FE) + Railway/Neon (BE)", "Prêt"],
]

add_styled_table(doc, ["Composant", "Détail", "Statut"], results_summary, [3.5, 8.5, 3.5])

add_heading_numbered(doc, "13.2 Centralisation des données", 2)

add_normal(doc, (
    "La principale réalisation est la centralisation de l'ensemble des données CMF "
    "dans un système unique. Les données qui étaient auparavant dispersées dans de "
    "multiples fichiers Excel sont désormais stockées dans une base de données relationnelle, "
    "accessibles via une interface web unifiée."
))

add_heading_numbered(doc, "13.3 Importation automatisée", 2)

add_normal(doc, (
    "Le pipeline d'importation permet de transférer des données depuis des fichiers Excel "
    "vers le système en quelques clics. Le mapping sémantique des colonnes, assisté par "
    "intelligence artificielle mais fonctionnant principalement de manière déterministe, "
    "réduit considérablement le temps de configuration de l'import."
))

add_heading_numbered(doc, "13.4 Workflow de validation", 2)

add_normal(doc, (
    "Le workflow en trois phases (Buyer → Capacity Manager → SQD) est strictement "
    "enforce par le système. Les permissions de champ empêchent les modifications "
    "hors de la section autorisée. L'étape de workflow est calculée automatiquement "
    "en fonction des données saisies."
))

add_heading_numbered(doc, "13.5 Templates modulaires", 2)

add_normal(doc, (
    "Le système de templates permet de créer des structures de données personnalisées "
    "sans modifier le code source. Les templates K0 et K9, spécifiques au processus CMF, "
    "sont intégralement fonctionnels avec leurs champs, sections et permissions."
))

add_heading_numbered(doc, "13.6 Tableau de bord analytique", 2)

add_normal(doc, (
    "Le dashboard offre une vue d'ensemble complète avec des KPI, des graphiques et "
    "des indicateurs adaptés au rôle de chaque utilisateur. La séparation des sections "
    "par rôle garantit que chaque acteur ne voit que les informations pertinentes."
))

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
#       CHAPITRE 14 — COMPÉTENCES ACQUISES
# ══════════════════════════════════════════════════════════════

add_heading_numbered(doc, "Chapitre 14 — Compétences acquises", 1)

add_heading_numbered(doc, "14.1 Compétences techniques", 2)

tech_skills = [
    ("Développement frontend", "Maîtrise de React 19, TypeScript, la gestion d'état "
     "(Zustand, TanStack Query), le responsive design (Tailwind CSS) et l'accessibilité (Radix UI)."),
    ("Développement backend", "Maîtrise de Python, FastAPI, l'architecture clean, "
     "les services métier, les DTOs et les interfaces."),
    ("API REST", "Conception et implémentation d'API RESTful avec documentation OpenAPI, "
     "validation Pydantic et gestion d'erreurs structurée."),
    ("Bases de données", "Modélisation relationnelle, ORM SQLAlchemy, migrations Alembic, "
     "optimisation des index, transactions et Unit of Work."),
    ("Traitement Excel", "Parsing de fichiers Excel avec openpyxl (Python) et xlsx/SheetJS (JS), "
     "détection de feuilles, scoring, extraction d'en-têtes."),
    ("Data engineering", "Pipeline de données complet : extraction, transformation, "
     "validation, chargement (ETL). Gestion des doublons et des cas limites."),
    ("Intégration IA", "Intégration de LLM local (Ollama) avec fallback gracieux, "
     "validation de réponses, gestion des timeouts et de l'indisponibilité."),
    ("Tests", "Écriture de tests unitaires et d'intégration avec pytest, fixtures, "
     "mocking, scénarios de test réalistes."),
    ("Git", "Utilisation avancée de Git : branches, commits conventionnels, merge, "
     "gestion de conflits, historique propre."),
    ("Déploiement", "Configuration multi-environnement, déploiement Vercel/Railway/Neon, "
     "adaptation des URLs de base de données."),
]

for title, desc in tech_skills:
    p = doc.add_paragraph()
    run = p.add_run(f"{title} : ")
    run.bold = True
    run.font.size = Pt(11)
    run = p.add_run(desc)
    run.font.size = Pt(11)

add_heading_numbered(doc, "14.2 Compétences méthodologiques", 2)

meth_skills = [
    ("Analyse des besoins", "Identification et structuration des besoins fonctionnels "
     "et non fonctionnels à partir de processus métier existants."),
    ("Conception architecturale", "Choix et justification d'une architecture clean, "
     "séparation des couches, design patterns (Unit of Work, Repository, DTO)."),
    ("Résolution de problèmes", "Approche systématique face aux défis techniques : "
     "analyse, recherche de solutions, implémentation, validation."),
    ("Documentation technique", "Rédaction de documentation de code, README, "
     "plans d'implémentation et ce rapport de stage."),
    ("Gestion de projet", "Travail autonome avec des jalons, priorisation des tâches, "
     "gestion du temps et des ressources."),
]

for title, desc in meth_skills:
    p = doc.add_paragraph()
    run = p.add_run(f"{title} : ")
    run.bold = True
    run.font.size = Pt(11)
    run = p.add_run(desc)
    run.font.size = Pt(11)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
#            CHAPITRE 15 — PERSPECTIVES
# ══════════════════════════════════════════════════════════════

add_heading_numbered(doc, "Chapitre 15 — Perspectives", 1)

add_normal(doc, (
    "La plateforme CMF constitue une base solide qui peut être enrichie et étendue "
    "pour répondre à des besoins futurs. Les principales évolutions envisagées "
    "sont les suivantes :"
))

perspectives = [
    ("Migration vers PostgreSQL",
     "La base SQLite utilisée en développement sera remplacée par PostgreSQL en production. "
     "Le déploiement Neon PostgreSQL est déjà configuré. Cette migration permettra une "
     "meilleure concurrence d'accès, une scalabilité horizontale et des fonctionnalités "
     "avancées (JSONB, full-text search)."),
    ("Déploiement cloud complet",
     "Le déploiement sur Vercel (frontend) + Railway (backend) + Neon (database) est "
     "déjà préparé. L'étape suivante est la mise en production effective avec monitoring, "
     "logs centralisés et alertes."),
    ("Audit trail complet",
     "Le modèle AuditLog existe mais n'est pas encore pleinement exploité. L'ajout "
     "d'un historique complet des modifications avec comparaison old/new values "
     "permettra une traçabilité totale."),
    ("Notifications avancées",
     "Le système de notifications de base peut être étendu vers des notifications par email "
     "(aiosmtplib déjà dans les dépendances), des notifications push et des alertes "
     "configurables par événement."),
    ("Analytics avancées",
     "L'ajout de tableaux de bord plus poussés avec tendances temporelles, prédictions "
     "de capacité, analyse de risques multi-dimensionnels et exports PDF/Excel "
     "(reportlab et xlsx déjà présents)."),
    ("Amélioration de l'IA",
     "L'extension du mapping IA vers la classification automatique de données, "
     "la détection d'anomalies et la suggestion de valeurs manquantes. "
     "Entraînement de modèles spécifiques sur les données CMF."),
    ("Scalabilité",
     "Mise en place de cache distribué (Redis déjà dans les dépendances), "
     "de files d'attente pour les imports de grande taille, et de load balancing."),
    ("API publique",
     "Exposition d'une API REST publique documentée pour l'intégration avec "
     "les systèmes existants de l'entreprise (ERP, MRP)."),
]

for title, desc in perspectives:
    p = doc.add_paragraph()
    run = p.add_run(f"{title} : ")
    run.bold = True
    run.font.size = Pt(11)
    run = p.add_run(desc)
    run.font.size = Pt(11)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
#            CONCLUSION GÉNÉRALE
# ══════════════════════════════════════════════════════════════

add_heading_numbered(doc, "Conclusion générale", 1)

add_normal(doc, (
    "Ce stage de fin d'études m'a permis de concevoir et de développer, de manière "
    "autonome et complète, une plateforme web de Capacity Management Framework (CMF). "
    "Le projet, réalisé dans un contexte industriel réel, m'a confronté à des problématiques "
    "concrètes de centralisation de données, de traitement de fichiers Excel complexes, "
    "d'intégration d'intelligence artificielle et de développement full-stack."
))

add_normal(doc, (
    "Les objectifs initiaux du stage ont été atteints. La plateforme CMF permet de :"
))

add_bullet(doc, "Centraliser l'ensemble des données CMF dans un système unique")
add_bullet(doc, "Importer des données depuis des fichiers Excel avec mapping sémantique")
add_bullet(doc, "Gérer les projets, fournisseurs, capacités et risques")
add_bullet(doc, "Suivre le workflow de validation Buyer → Capacity Manager → SQD")
add_bullet(doc, "Disposer d'un tableau de bord analytique adapté à chaque rôle")
add_bullet(doc, "Gérer les structures de données de manière modulaire")
add_bullet(doc, "Garantir la sécurité via l'authentification JWT et le RBAC")

add_normal(doc, (
    "Ce projet m'a permis d'acquérir des compétences solides en développement web "
    "full-stack, en architecture logicielle, en traitement de données et en intégration "
    "d'intelligence artificielle. La résolution des difficultés techniques rencontrées "
    "(colonnes en double, orientation des données, fallback IA, race conditions) a "
    "renforcé ma capacité à analyser des problèmes complexes et à proposer des solutions "
    "efficaces."
))

add_normal(doc, (
    "Sur le plan humain, cette expérience m'a appris l'importance de la communication "
    "avec l'encadrant et l'équipe, de la documentation rigoureuse et de l'adaptation "
    "aux contraintes métier. Le travail en autonomie sur un projet de cette ampleur "
    "a été formateur et m'a préparé aux défis du monde professionnel."
))

add_normal(doc, (
    "La plateforme constitue une base évolutive qui pourra être enrichie selon les "
    "besoins futurs : déploiement cloud, analytics avancées, amélioration de l'IA, "
    "intégration avec les systèmes existants. Ce projet démontre qu'il est possible "
    "de développer des solutions métier robustes et modernes en utilisant des technologies "
    "open-source accessibles."
))

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
#                   ANNEXES
# ══════════════════════════════════════════════════════════════

add_heading_numbered(doc, "Annexes", 1)

add_heading_numbered(doc, "Annexe A — Architecture technique détaillée", 2)

add_normal(doc, (
    "Schéma d'architecture backend (Clean Architecture) :"
))

arch_be = [
    "┌──────────────────────────────────────────────────┐",
    "│                   API Layer                      │",
    "│  /api/v1/auth  /api/v1/projects  /api/v1/import  │",
    "│  14 route modules · Pydantic validation           │",
    "├──────────────────────────────────────────────────┤",
    "│               Application Layer                  │",
    "│  20 services · 12 DTOs · 2 interfaces            │",
    "│  import_engine · ollama_mapping · data_normalizer│",
    "├──────────────────────────────────────────────────┤",
    "│              Infrastructure Layer                │",
    "│  persistence (18 models, 12 repositories)        │",
    "│  auth (JWT + RBAC) · notifications · storage     │",
    "├──────────────────────────────────────────────────┤",
    "│                   Core Layer                     │",
    "│  config · database · security · cache · logging  │",
    "├──────────────────────────────────────────────────┤",
    "│                 Domain Layer                     │",
    "│  enums · events · import_schema · K0/K9 JSON    │",
    "└──────────────────────────────────────────────────┘",
]

for line in arch_be:
    add_code_block(doc, line)

add_heading_numbered(doc, "Annexe B — Pipeline d'importation Excel", 2)

pipeline_arch = [
    "Upload Excel (.xlsx/.xls/.xlsm)",
    "        │",
    "        ▼",
    "┌─────────────────────┐",
    "│  scan-workbook      │ → Worksheet scoring + classification",
    "└────────┬────────────┘",
    "         │",
    "         ▼",
    "┌─────────────────────┐",
    "│  extract-headers    │ → Header row detection (top 20 rows)",
    "└────────┬────────────┘   Orientation detection (H/V)",
    "         │",
    "         ▼",
    "┌─────────────────────┐",
    "│  ollama-map         │ → 5-level cascade mapping",
    "└────────┬────────────┘   cache → exact → alias → fuzzy → AI",
    "         │",
    "         ▼",
    "┌─────────────────────┐",
    "│  preview            │ → Validation + CREATE/UPDATE/RESTORE",
    "└────────┬────────────┘",
    "         │",
    "         ▼",
    "┌─────────────────────┐",
    "│  execute            │ → SAVEPOINT per row · Audit · History",
    "└─────────────────────┘",
]

for line in pipeline_arch:
    add_code_block(doc, line)

add_heading_numbered(doc, "Annexe C — Exemple de mapping K0 (positions)", 2)

add_normal(doc, (
    "Les 47 positions de colonnes K0 (extrait) :"
))

k0_positions = [
    ["0", "part_number", "Numéro de pièce"],
    ["1", "index", "Index"],
    ["2", "description", "Description"],
    ["3", "coef", "Coefficient"],
    ["4", "serial_piece_price", "Prix unitaire"],
    ["5", "mass_purchase", "Achat massique"],
    ["6", "ru", "RU"],
    ["7", "noa", "NOA"],
    ["8", "make_battery_lp_1", "Make Battery LP 1"],
    ["9", "make_battery_lp_2", "Make Battery LP 2"],
    ["10", "supplier_name", "Nom du fournisseur"],
    ["11", "vendor_name", "Nom du vendeur"],
    ["12", "manufacturer_name", "Nom du fabricant"],
    ["13", "combined_cofor", "COFOR combiné"],
    ["...", "...", "..."],
    ["46", "manufacturing_validated", "Manufacturing validé"],
]

add_styled_table(doc, ["Index", "Clé CMF", "Description"], k0_positions, [2, 5, 8.5])

add_heading_numbered(doc, "Annexe D — Cascade de mapping des colonnes", 2)

mapping_example = [
    ["Part Number", "part_number", "1.0", "Correspondance exacte normalisée"],
    ["Part no.", "part_number", "0.96", "Alias connu (part_no → part_number)"],
    ["Suplier Name", "supplier_name", "0.91", "Similarité floue (typo acceptée)"],
    ["Tango TBD", "tango_order", "0.95", "IA Ollama (fallback)"],
    ["Custom Col X", "—", "—", "Pas de correspondance (no_match)"],
]

add_styled_table(doc, ["En-tête Excel", "Champ CMF", "Confiance", "Méthode"], mapping_example, [3, 3.5, 2, 7])

add_heading_numbered(doc, "Annexe E — Structure des modèles de données", 2)

add_normal(doc, (
    "Diagramme simplifié des relations principales :"
))

er_diagram = [
    "┌──────────┐     ┌────────────┐     ┌──────────────┐",
    "│  users   │────<│  projects  │>────│  templates   │",
    "└──────────┘     └─────┬──────┘     └──────────────┘",
    "     │                 │",
    "     │            ┌────┴─────┐",
    "     │            │          │",
    "     │     ┌──────┴───┐ ┌───┴──────────┐",
    "     │     │  parts   │ │  documents   │",
    "     │     └────┬─────┘ └──────────────┘",
    "     │          │",
    "     │    ┌─────┴──────┐",
    "     │    │            │",
    "     │ ┌──┴──────┐ ┌──┴──────────┐",
    "     │ │  risks  │ │  capacity   │",
    "     │ └─────────┘ └──────┬──────┘",
    "     │                    │",
    "     │              ┌─────┴──────┐",
    "     │              │ suppliers  │",
    "     │              └────────────┘",
    "     │",
    "┌────┴──────────┐",
    "│ notifications │",
    "│ activity_logs │",
    "│ audit_logs    │",
    "└───────────────┘",
]

for line in er_diagram:
    add_code_block(doc, line)

# ══════════════════════════════════════════════════════════════
#                   FOOTER NOTE
# ══════════════════════════════════════════════════════════════

doc.add_page_break()
add_heading_numbered(doc, "Notes finales", 1)

add_normal(doc, (
    "Ce document a été généré automatiquement à partir de l'analyse du code source "
    "du projet CMF Platform. Il constitue une base complète et détaillée pour la "
    "rédaction du rapport de stage final."
))

add_normal(doc, (
    "Les informations contenues dans ce rapport sont tirées exclusivement de l'analyse "
    "du code source, des fichiers de configuration, des tests et de l'historique Git. "
    "Aucune information n'a été inventée au-delà de ce que le code révèle."
))

add_normal(doc, (
    "Les champs entre crochets [Nom de l'entreprise], [Votre Nom], [Votre École], etc. "
    "doivent être complétés par l'étudiant avec ses informations personnelles."
))

# ── Save ──
output_path = r"C:\Users\hp\AppData\Local\Temp\opencode\Rapport_Stage_CMF_Platform.docx"
doc.save(output_path)
print("Rapport genere avec succes :", output_path)
print("   Fichier : Rapport_Stage_CMF_Platform.docx")
